"""
Vellum-compatible DOCX export.

Vellum's import rules (from research):
  - Chapter title must be Word's Heading 1 style
  - Each chapter should start with a page break, then Heading 1
  - Scene breaks: single blank line OR centered *** (or ###, +++)
  - No manual indentation in paragraphs
  - Use center+bold for "Chapter N" title as an alternative signal
  - Subheads use Heading 2 / Heading 3

This module generates a .docx that drops into Vellum with chapters and
scene breaks auto-detected.

Backed by python-docx (already installed).
"""
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_page_break(paragraph):
    """Insert a hard page break into a paragraph."""
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def _set_style_paragraph(paragraph, style_name):
    """Set a paragraph's style by name."""
    try:
        paragraph.style = style_name
    except KeyError:
        pass


def build_vellum_docx(
    project_title: str,
    author: str,
    chapters: list[dict],
    front_matter: dict | None = None,
    dedication: str = "",
    epigraph: str = "",
    style_notes: str = "",
) -> bytes:
    """Build a Vellum-compatible DOCX from chapter dicts.

    chapters: list of {
        "number": int,
        "title": str,
        "content": str,  # the chapter's markdown body
        "scenes": [{"name": str, "content": str}] (optional)
    }

    Returns the .docx as bytes.
    """
    front_matter = front_matter or {}
    doc = Document()

    # Configure default paragraph style to have no first-line indent
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.first_line_indent = Inches(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5

    # Set Heading 1 — chapter title
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Georgia"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Inches(0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.page_break_before = True  # <-- key for Vellum

    # Set Heading 2 — scene subheads
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Georgia"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Inches(0)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    # Title page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    run = title_p.add_run(project_title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = "Georgia"

    if author:
        author_p = doc.add_paragraph()
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_p.paragraph_format.space_before = Pt(24)
        run = author_p.add_run(f"by {author}")
        run.font.size = Pt(14)
        run.italic = True
        run.font.name = "Georgia"

    # Dedication (if any)
    if dedication:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(200)
        run = p.add_run(dedication)
        run.italic = True
        run.font.size = Pt(12)

    # Epigraph (if any)
    if epigraph:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not dedication:
            p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(60)
        run = p.add_run(epigraph)
        run.italic = True
        run.font.size = Pt(11)

    # Style notes (if any)
    if style_notes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.page_break_before = True
        run = p.add_run(f"Style: {style_notes}")
        run.italic = True
        run.font.size = Pt(10)

    # Chapters
    for ch in chapters:
        # Chapter heading — Heading 1 style triggers Vellum's chapter detection
        ch_title = ch.get("title", f"Chapter {ch.get('number', '?')}")
        ch_number = ch.get("number", "")
        # Combine: "Chapter 1: The Beginning" — Vellum parses this format
        if ch_number and not re.match(rf"^(chapter|prologue|epilogue|part)\b",
                                       ch_title, re.IGNORECASE):
            heading_text = f"Chapter {ch_number}: {ch_title}"
        else:
            heading_text = ch_title

        # Add the heading. page_break_before is set in the style, so the
        # paragraph automatically starts on a new page.
        h = doc.add_paragraph(style="Heading 1")
        h.add_run(heading_text)

        # Body — strip any leading H1 from the chapter content (it duplicates
        # our heading above) and convert markdown to Vellum-friendly paragraphs
        body = _strip_leading_heading(ch.get("content", ""), level=1)
        _emit_chapter_body(doc, body)

        # Scenes (if present) — sub-chapters become Heading 2 subsections
        for scene in ch.get("scenes", []):
            sub = doc.add_paragraph(style="Heading 2")
            sub.add_run(scene.get("name", "Scene"))
            _emit_chapter_body(doc, scene.get("content", ""))

    # Save to bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _emit_chapter_body(doc, content: str):
    """Emit a chapter's body as a sequence of paragraphs.

    Rules:
      - # / ## / ### headings (markdown) become Heading 1/2/3
      - A line with only *** / ### / +++ (or a single blank line) becomes
        a scene break (centered, with the symbol preserved)
      - All other lines become Normal paragraphs
      - Inline **bold** and *italic* are converted
    """
    lines = content.split("\n")
    consecutive_blanks = 0
    for raw in lines:
        line = raw.rstrip()

        # Blank line — count toward scene break
        if not line.strip():
            consecutive_blanks += 1
            continue

        # Scene break marker (***, ###, +++) on its own line
        if re.match(r"^\s*(\*\*\*|###|\+\+\+)\s*$", line):
            _add_scene_break(doc, line.strip())
            consecutive_blanks = 0
            continue

        # Markdown heading
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            _emit_inline(p, line[4:])
            consecutive_blanks = 0
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            _emit_inline(p, line[3:])
            consecutive_blanks = 0
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            _emit_inline(p, line[2:])
            consecutive_blanks = 0
            continue

        # Regular paragraph
        # If we had ≥1 blank lines before, add a scene break (Vellum reads
        # a single blank line as a scene break).
        if consecutive_blanks >= 1:
            _add_scene_break(doc, "* * *")
        p = doc.add_paragraph()
        _emit_inline(p, line)
        consecutive_blanks = 0


def _add_scene_break(doc, marker: str = "* * *"):
    """Add a centered scene break line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(marker)
    run.font.size = Pt(11)


def _strip_leading_heading(content: str, level: int = 1) -> str:
    """Remove a leading markdown heading of the given level from content.

    We strip leading blank lines, then if the first non-blank line is
    a level-# heading, remove that line. Everything else stays.
    """
    lines = content.split("\n")
    out = []
    found_heading = False
    prefix = "#" * level + " "
    for line in lines:
        if not found_heading and not line.strip():
            out.append(line)
            continue
        if not found_heading and line.startswith(prefix):
            found_heading = True
            continue
        out.append(line)
    # Strip leading blank lines from the result
    while out and not out[0].strip():
        out.pop(0)
    return "\n".join(out)


def _emit_inline(paragraph, text: str):
    """Add text to a paragraph, handling **bold** and *italic* markers.

    Vellum preserves bold and italic, so we mark them in the runs.
    """
    # Parse **bold** and *italic* segments
    pos = 0
    pattern = re.compile(r"(\*\*([^*]+)\*\*|\*([^*]+)\*)")
    for m in pattern.finditer(text):
        # Pre-text
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(2):
            # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
        else:
            # italic
            run = paragraph.add_run(m.group(3))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# --------------------------------------------------------------------------
# Standard DOCX export (not Vellum-specific, used for general Word use)
# --------------------------------------------------------------------------

def build_standard_docx(
    project_title: str,
    author: str,
    chapters: list[dict],
    dedication: str = "",
    epigraph: str = "",
    style_notes: str = "",
) -> bytes:
    """Build a standard .docx with chapter titles as Heading 1 and prose.

    Different from Vellum version: no page_break_before, no scene-break
    conventions, just clean Word output.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(20)
    h1.font.bold = True

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(project_title)
    run.bold = True
    run.font.size = Pt(28)
    if author:
        a = doc.add_paragraph()
        a.alignment = WD_ALIGN_PARAGRAPH.CENTER
        a.add_run(f"by {author}").italic = True

    if dedication:
        d = doc.add_paragraph()
        d.add_run(f"Dedication: {dedication}").italic = True
    if epigraph:
        e = doc.add_paragraph()
        e.add_run(f"Epigraph: {epigraph}").italic = True
    if style_notes:
        s = doc.add_paragraph()
        s.add_run(f"Style: {style_notes}").italic = True

    for ch in chapters:
        # Page break + Heading 1
        p = doc.add_paragraph(style="Heading 1")
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        ch_title = ch.get("title", f"Chapter {ch.get('number', '?')}")
        ch_number = ch.get("number", "")
        if ch_number and not re.match(r"^(chapter|prologue|epilogue|part)\b",
                                       ch_title, re.IGNORECASE):
            p.add_run(f"Chapter {ch_number}: {ch_title}")
        else:
            p.add_run(ch_title)
        for line in (ch.get("content", "")).split("\n"):
            if line.strip():
                doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
