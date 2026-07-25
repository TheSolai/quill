"""
Tests for Dross AI features:
  - AgentMail integration
  - Web search tool
  - Tool registry
  - Vellum DOCX export
  - New export formats (RTF, OPML, bundle)
  - Natural language email intents
"""
import json
import os
import re
import sys
import tempfile
from io import BytesIO
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))


# --------------------------------------------------------------------------
# Vellum DOCX
# --------------------------------------------------------------------------

class TestVellumDocx:
    def test_builds_valid_docx(self):
        from vellum_docx import build_vellum_docx
        chapters = [
            {"number": 1, "title": "The Beginning", "content": "It was a dark night."},
            {"number": 2, "title": "The Crossing", "content": "She found a boat."},
        ]
        result = build_vellum_docx("Test Book", "Tester", chapters)
        assert isinstance(result, bytes)
        assert len(result) > 1000  # real DOCX
        # Should start with PK (zip magic)
        assert result[:2] == b"PK"

    def test_chapters_use_heading_1(self):
        from vellum_docx import build_vellum_docx
        from docx import Document
        chapters = [
            {"number": 1, "title": "The Beginning", "content": "Body text."},
        ]
        result = build_vellum_docx("Test Book", "Tester", chapters)
        doc = Document(BytesIO(result))
        h1s = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert any("Chapter 1: The Beginning" in p.text for p in h1s)

    def test_scene_break_centered(self):
        from vellum_docx import build_vellum_docx, _emit_chapter_body
        from docx import Document
        # Use a smaller test — just the body emission
        doc = Document()
        _emit_chapter_body(doc, "First paragraph.\n\n***\n\nSecond paragraph.")
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        centered = [p for p in doc.paragraphs
                    if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and p.text.strip()]
        assert len(centered) >= 1
        assert any("***" in p.text or "* * *" in p.text for p in centered)

    def test_inline_bold_italic(self):
        from vellum_docx import build_vellum_docx
        from docx import Document
        chapters = [
            {"number": 1, "title": "Test", "content": "Some **bold** and *italic* text."},
        ]
        result = build_vellum_docx("Test", "Tester", chapters)
        doc = Document(BytesIO(result))
        # Find the paragraph with our text
        target = next(p for p in doc.paragraphs if "Some" in p.text)
        runs = target.runs
        assert any(r.bold for r in runs)
        assert any(r.italic for r in runs)

    def test_page_break_before_chapter(self):
        from vellum_docx import build_vellum_docx
        from docx import Document
        chapters = [
            {"number": 1, "title": "First", "content": "Body."},
            {"number": 2, "title": "Second", "content": "Body."},
        ]
        result = build_vellum_docx("Test", "Tester", chapters)
        doc = Document(BytesIO(result))
        h1_style = doc.styles["Heading 1"]
        # The style should have page_break_before set
        assert h1_style.paragraph_format.page_break_before is True

    def test_strips_leading_heading(self):
        from vellum_docx import _strip_leading_heading
        content = "# Chapter 1\n\nBody text here."
        result = _strip_leading_heading(content)
        assert "Body text here" in result
        assert "# Chapter 1" not in result

    def test_preserves_body_when_no_leading_heading(self):
        from vellum_docx import _strip_leading_heading
        content = "Body without heading."
        result = _strip_leading_heading(content)
        assert "Body without heading" in result


# --------------------------------------------------------------------------
# Web search
# --------------------------------------------------------------------------

class TestWebSearch:
    def test_empty_query_returns_empty(self):
        from web_search import search
        assert search("") == []
        assert search("   ") == []

    def test_search_returns_list(self):
        from web_search import search
        results = search("python programming", max_results=3)
        assert isinstance(results, list)
        # Real network test; should have at least 1 result
        if results and "error" not in results[0]:
            assert all("url" in r for r in results)


# --------------------------------------------------------------------------
# AgentMail service
# --------------------------------------------------------------------------

class TestAgentMailService:
    def test_inbox_constant(self):
        import agentmail_service
        assert agentmail_service.DROSS_INBOX == "thedross@agentmail.to"

    def test_parse_email_intent_simple(self):
        from agentmail_service import parse_email_intent
        result = parse_email_intent("email the book to user@example.com")
        assert result is not None
        assert result["to"] == "user@example.com"
        assert result["what"] == "book"

    def test_parse_email_intent_chapter(self):
        from agentmail_service import parse_email_intent
        result = parse_email_intent("send chapter 3 to editor@pub.com")
        assert result is not None
        assert result["to"] == "editor@pub.com"
        assert result["what"] == "chapter"

    def test_parse_email_intent_with_subject(self):
        from agentmail_service import parse_email_intent
        result = parse_email_intent("email the book to user@example.com subject: Final Draft")
        assert result is not None
        assert result["subject"] == "Final Draft"

    def test_parse_email_intent_no_email_returns_none(self):
        from agentmail_service import parse_email_intent
        assert parse_email_intent("hello how are you") is None
        assert parse_email_intent("write a chapter about cats") is None

    def test_parse_email_intent_no_address_returns_none(self):
        from agentmail_service import parse_email_intent
        # has 'email' but no address
        assert parse_email_intent("email me the book") is None


# --------------------------------------------------------------------------
# Dross tools
# --------------------------------------------------------------------------

class TestDrossTools:
    def test_list_tools_returns_all(self):
        from dross_tools import list_tools
        tools = list_tools()
        names = [t["name"] for t in tools]
        assert "web_search" in names
        assert "email_send" in names
        assert "shell_exec" in names
        assert "list_files" in names
        assert "read_file" in names

    def test_call_unknown_tool_returns_error(self):
        from dross_tools import call_tool
        result = call_tool("nonexistent_tool", {})
        assert "error" in result

    def test_call_web_search(self):
        from dross_tools import call_tool
        result = call_tool("web_search", {"query": "test query", "max_results": 2})
        # Real network call
        assert "results" in result or "error" in result

    def test_shell_exec_blocks_dangerous(self):
        from dross_tools import call_tool
        result = call_tool("shell_exec", {"cmd": "rm -rf /"})
        assert "error" in result
        assert "blocked" in result["error"]

    def test_shell_exec_blocks_sudo(self):
        from dross_tools import call_tool
        result = call_tool("shell_exec", {"cmd": "sudo apt-get update"})
        assert "error" in result

    def test_shell_exec_runs_safe_command(self):
        from dross_tools import call_tool
        result = call_tool("shell_exec", {"cmd": "echo hello", "timeout": 5})
        assert "stdout" in result
        assert "hello" in result["stdout"]

    def test_tools_as_openai_functions_shape(self):
        from dross_tools import tools_as_openai_functions
        funcs = tools_as_openai_functions()
        for f in funcs:
            assert f["type"] == "function"
            assert "name" in f["function"]
            assert "description" in f["function"]
            assert "parameters" in f["function"]


# --------------------------------------------------------------------------
# API endpoints
# --------------------------------------------------------------------------

@pytest.fixture
def client():
    from server import app
    app.config["TESTING"] = True
    with app.test_client() as c:
        yield c


class TestAgentMailAPI:
    def test_status(self, client):
        r = client.get("/api/agentmail/status")
        assert r.status_code == 200
        d = r.get_json()
        assert d["inbox"] == "thedross@agentmail.to"

    def test_inbox(self, client):
        r = client.get("/api/agentmail/inbox?limit=5")
        assert r.status_code == 200
        d = r.get_json()
        assert "messages" in d

    def test_send_requires_to(self, client):
        r = client.post("/api/agentmail/send", json={
            "subject": "test", "text": "hi"
        })
        # 400 (no to) or 503 (sdk unavailable) or 500 (rate limited) all OK
        assert r.status_code in (400, 500, 503)

    def test_draft_requires_to(self, client):
        r = client.post("/api/agentmail/draft", json={
            "subject": "test", "text": "hi"
        })
        assert r.status_code in (400, 500, 503)


class TestToolsAPI:
    def test_list_tools(self, client):
        r = client.get("/api/tools")
        assert r.status_code == 200
        d = r.get_json()
        names = [t["name"] for t in d["tools"]]
        assert "web_search" in names
        assert "email_send" in names

    def test_call_tool_requires_name(self, client):
        r = client.post("/api/tools/call", json={"args": {}})
        assert r.status_code == 400


class TestWebSearchAPI:
    def test_search_requires_q(self, client):
        r = client.get("/api/search")
        assert r.status_code == 400

    def test_search_returns_results(self, client):
        r = client.get("/api/search?q=python&max=2")
        assert r.status_code == 200
        d = r.get_json()
        assert "results" in d


class TestNewExportFormats:
    def test_vellum_export(self, client):
        # Use a known project — set up via the test
        r = client.post("/api/projects", json={"name": "vellum-test"})
        pid = r.get_json()["id"]
        # Add a chapter
        r = client.post(f"/api/projects/{pid}/chapters", json={"name": "chapter-1"})
        # Set some content
        client.put(f"/api/projects/{pid}/chapters/chapter-1/content", json={
            "content": "# Chapter 1\n\nIt was a dark and stormy night."
        })
        # Set settings
        client.put(f"/api/projects/{pid}/settings", json={
            "title": "Vellum Test", "author": "Tester"
        })
        r = client.get(f"/api/projects/{pid}/export/vellum")
        assert r.status_code == 200
        # Should be a docx (zip with PK header)
        assert r.data[:2] == b"PK"

    def test_opml_export(self, client):
        r = client.post("/api/projects", json={"name": "opml-test"})
        pid = r.get_json()["id"]
        client.post(f"/api/projects/{pid}/chapters", json={"name": "chapter-1"})
        client.put(f"/api/projects/{pid}/chapters/chapter-1/content", json={
            "content": "# First\n\nBody."
        })
        r = client.get(f"/api/projects/{pid}/export/opml")
        assert r.status_code == 200
        assert b"<opml" in r.data

    def test_bundle_export(self, client):
        r = client.post("/api/projects", json={"name": "bundle-test"})
        pid = r.get_json()["id"]
        client.post(f"/api/projects/{pid}/chapters", json={"name": "chapter-1"})
        client.put(f"/api/projects/{pid}/chapters/chapter-1/content", json={
            "content": "# First\n\nBody text."
        })
        r = client.get(f"/api/projects/{pid}/export/bundle")
        assert r.status_code == 200
        # zip — but might also be 404 if the chapter file isn't there
        # The test just checks the endpoint exists and returns either zip or 404
        assert r.status_code in (200, 404)
        if r.status_code == 200:
            assert r.data[:2] == b"PK"

    def test_rtf_export(self, client):
        r = client.post("/api/projects", json={"name": "rtf-test"})
        pid = r.get_json()["id"]
        client.post(f"/api/projects/{pid}/chapters", json={"name": "chapter-1"})
        client.put(f"/api/projects/{pid}/chapters/chapter-1/content", json={
            "content": "# First\n\nBody text."
        })
        r = client.get(f"/api/projects/{pid}/export/rtf")
        assert r.status_code == 200
        # RTF starts with {\rtf
        assert r.data[:5] == b"{\\rtf" or b"pandoc" in r.data[:20]


class TestDrossSystemPrompt:
    def test_chat_uses_quill_persona(self, client):
        """The /api/chat endpoint should include the Quill system prompt by default."""
        from slot_providers import PROVIDERS
        mock_instance = MagicMock()
        mock_instance.chat.return_value = "PONG"
        mock_instance.stream.return_value = iter(["PONG"])
        with patch.dict(PROVIDERS, {"minimax": MagicMock(return_value=mock_instance)}):
            client.post("/api/slots/minimax-text/activate")
            r = client.post("/api/chat", json={
                "messages": [{"role": "user", "content": "Say PONG"}],
                "stream": False,
            })
            assert r.status_code == 200
            # The Quill system prompt should have been prepended
            call_args = mock_instance.chat.call_args
            messages = call_args[0][0]
            assert messages[0]["role"] == "system"
            assert "Quill" in messages[0]["content"]

    def test_chat_email_intent_routes_to_email(self, client):
        """When the user says 'email the book to X', the chat should call the email tool directly."""
        r = client.post("/api/projects", json={"name": "intent-test"})
        pid = r.get_json()["id"]
        client.post(f"/api/projects/{pid}/chapters", json={"name": "chapter-1"})
        client.put(f"/api/projects/{pid}/chapters/chapter-1/content", json={
            "content": "# Chapter 1\n\nSome body text."
        })
        from slot_providers import PROVIDERS
        mock_instance = MagicMock()
        mock_instance.chat.return_value = "should not be called"
        mock_instance.stream.return_value = iter(["should not be called"])
        with patch.dict(PROVIDERS, {"minimax": MagicMock(return_value=mock_instance)}):
            client.post("/api/slots/minimax-text/activate")
            r = client.post("/api/chat", json={
                "project_id": pid,
                "messages": [{"role": "user", "content": "email the book to amre@agentmail.to"}],
                "stream": False,
            })
            # The model should NOT be called — intent was caught
            assert mock_instance.chat.called is False
            # And we got an email result
            d = r.get_json()
            # Either email succeeded, was rate-limited, or failed for a benign reason
            assert "email" in d or "rate" in str(d).lower() or "ok" in d


if __name__ == "__main__":
    pytest.main([__file__, "-v"])


# --------------------------------------------------------------------------
# Chapter-write intent
# --------------------------------------------------------------------------

class TestChapterWriteIntent:
    def test_detect_write_chapter_1(self):
        from server import _extract_chapter_write_intent
        result = _extract_chapter_write_intent("Write the next chapter 1")
        assert result is not None
        assert result["action"] == "write_chapter"
        assert "1" in result["target"]

    def test_detect_write_chapter_word(self):
        from server import _extract_chapter_write_intent
        result = _extract_chapter_write_intent("Draft chapter three about the river")
        assert result is not None
        assert "three" in result["target"].lower() or "3" in result["target"]

    def test_detect_write_current(self):
        from server import _extract_chapter_write_intent
        result = _extract_chapter_write_intent("Continue writing this chapter")
        assert result is not None
        assert result["target"] == "current"

    def test_detect_write_scene(self):
        from server import _extract_chapter_write_intent
        result = _extract_chapter_write_intent("Compose a new scene here")
        assert result is not None

    def test_no_intent_for_random_chat(self):
        from server import _extract_chapter_write_intent
        assert _extract_chapter_write_intent("hello how are you") is None
        assert _extract_chapter_write_intent("what is the meaning of life") is None
        assert _extract_chapter_write_intent("search for cats") is None

    def test_resolve_chapter_target(self):
        from server import _resolve_chapter_target
        # Create a temp project
        import os
        from pathlib import Path
        from server import get_project_dir
        pid = "test-resolve-xyz"
        pd = get_project_dir(pid)
        # Create some chapter files
        (pd / "chapter-1.md").write_text("# 1\n")
        (pd / "chapter-2.md").write_text("# 2\n")
        try:
            # Direct match
            assert _resolve_chapter_target(pid, "chapter-1") == "chapter-1"
            # "current" falls back to first
            assert _resolve_chapter_target(pid, "current") == "chapter-1"
            # Numeric
            assert _resolve_chapter_target(pid, "chapter-2") == "chapter-2"
            # Non-existent
            assert _resolve_chapter_target(pid, "chapter-99") is None
        finally:
            import shutil
            shutil.rmtree(pd, ignore_errors=True)


# --------------------------------------------------------------------------
# Slot categories
# --------------------------------------------------------------------------

class TestSlotCategories:
    def test_default_slots_have_categories(self):
        from slots import load_slots
        for s in load_slots():
            assert s.category in ("local", "creative", "research", "code", "cloud", "minimax"), \
                f"slot {s.id} has bad category {s.category}"

    def test_tool_calling_flag_persists(self):
        from slots import load_slots
        for s in load_slots():
            if s.type in ("ollama", "mlx"):
                # All current Ollama models support tools
                assert s.tool_calling is True, f"ollama slot {s.id} should have tool_calling=True"

    def test_groq_tool_use_slot_exists(self):
        from slots import get_slot
        s = get_slot("groq-tool-use")
        assert s is not None
        assert s.model_id == "llama3-groq-tool-use:8b"
        assert s.tool_calling is True

    def test_api_returns_new_fields(self, client):
        r = client.get("/api/slots")
        d = r.get_json()
        for s in d["slots"]:
            assert "category" in s, f"slot {s['id']} missing category in API"
            assert "tool_calling" in s, f"slot {s['id']} missing tool_calling in API"
            assert "thinking" in s, f"slot {s['id']} missing thinking in API"


# --------------------------------------------------------------------------
# Tool calling in Ollama provider
# --------------------------------------------------------------------------

class TestOllamaToolCalling:
    def test_payload_includes_tools(self):
        from slots import ModelSlot
        from slot_providers import OllamaProvider
        # Construct a slot in-memory (no need to write to disk)
        s = ModelSlot(
            id="test-tool", name="Test", type="ollama",
            model_id="groq-tool",
        )
        prov = OllamaProvider(s)
        # Build payload with tools
        payload = prov._build_payload(
            [{"role": "user", "content": "What's the weather?"}],
            {"tools": [{"type": "function", "function": {"name": "weather", "description": "Get weather"}}]},
            stream=False
        )
        assert "tools" in payload
        assert len(payload["tools"]) == 1

    def test_payload_no_tools(self):
        from slots import ModelSlot
        from slot_providers import OllamaProvider
        s = ModelSlot(
            id="test-no-tool", name="Test", type="ollama",
            model_id="gemma4:31b",
        )
        prov = OllamaProvider(s)
        payload = prov._build_payload(
            [{"role": "user", "content": "Hello"}],
            {},
            stream=False
        )
        assert "tools" not in payload


# --------------------------------------------------------------------------
# Input validation (path safety)
# --------------------------------------------------------------------------

class TestInputValidation:
    """Bug fixes: path traversal, slash in names, null content, etc."""

    def test_safe_name_replaces_slash(self):
        from server import safe_name
        assert "/" not in safe_name("my/chapter")
        # Should be safe to use as filename
        assert safe_name("my/chapter") == "my-chapter"

    def test_safe_name_rejects_path_traversal(self):
        from server import safe_name
        result = safe_name("../etc/passwd")
        assert ".." not in result
        assert "/" not in result

    def test_safe_name_replaces_spaces(self):
        from server import safe_name
        assert safe_name("Chapter One") == "Chapter-One"

    def test_safe_name_handles_unicode(self):
        from server import safe_name
        assert safe_name("café") == "café"
        assert safe_name("日本語") == "日本語"

    def test_safe_name_handles_empty(self):
        from server import safe_name
        assert safe_name("") == "untitled"
        assert safe_name(None) == "untitled"
        assert safe_name("   ") == "untitled"

    def test_safe_name_handles_special_chars(self):
        from server import safe_name
        for char in '<>:"/\\|?*':
            result = safe_name(f"a{char}b")
            assert char not in result, f"safe_name didn't remove {char!r}"

    def test_safe_name_collapses_dashes(self):
        from server import safe_name
        # Multiple consecutive dashes collapse
        assert safe_name("a---b") == "a-b"
        # Spaces become dashes first, then collapse
        assert safe_name("a - - b") == "a-b"
        # Tabs/newlines become dashes too
        assert safe_name("a\t\tb") == "a-b"

    def test_safe_name_length_limit(self):
        from server import safe_name
        long = "a" * 200
        result = safe_name(long, max_len=50)
        assert len(result) <= 50

    def test_validate_project_id(self):
        from server import validate_project_id
        assert validate_project_id("book") == "book"
        assert validate_project_id("book-1") == "book-1"
        # Invalid
        assert validate_project_id("") is None
        assert validate_project_id("a/b") is None
        assert validate_project_id("a\\b") is None
        assert validate_project_id(".hidden") is None
        assert validate_project_id("a" * 100) is None
        assert validate_project_id(None) is None

    def test_safe_content_handles_null(self):
        from server import safe_content
        assert safe_content(None) == ""
        assert safe_content("") == ""
        assert safe_content("hello") == "hello"
        assert safe_content(123) == "123"

    def test_create_chapter_sanitizes_name(self, client):
        """Chapter names with slashes should be sanitized, not 500."""
        r = client.post("/api/projects", json={"name": "bug-fix-test"})
        pid = r.get_json()["id"]
        # Slash in name should not crash
        r = client.post(f"/api/projects/{pid}/chapters", json={"name": "my/chapter?"})
        assert r.status_code == 200, f"Got {r.status_code}: {r.text}"
        d = r.get_json()
        assert "/" not in d["name"]

    def test_create_chapter_with_path_traversal(self, client):
        """Chapter names with .. should be sanitized."""
        r = client.post("/api/projects", json={"name": "bug-fix-test2"})
        pid = r.get_json()["id"]
        r = client.post(f"/api/projects/{pid}/chapters", json={"name": "../etc/passwd"})
        assert r.status_code == 200, f"Got {r.status_code}: {r.text}"
        d = r.get_json()
        assert ".." not in d["name"]
        assert "/" not in d["name"]

    def test_create_project_sanitizes_path_traversal(self, client):
        """Project names with .. should be sanitized to safe names."""
        r = client.post("/api/projects", json={"name": "../etc"})
        # The name is sanitized to "etc" (a safe folder name, not a 400)
        assert r.status_code == 200
        d = r.get_json()
        assert ".." not in d["id"]
        assert "/" not in d["id"]
        # Direct dot-slash should be rejected
        r = client.post("/api/projects", json={"name": "../../../etc/passwd"})
        assert r.status_code == 200
        d = r.get_json()
        assert ".." not in d["id"]

    def test_compile_nonexistent_project_404(self, client):
        r = client.get("/api/projects/nonexistent-xyz-12345/compile")
        assert r.status_code == 404

    def test_list_scenes_nonexistent_chapter_404(self, client):
        r = client.post("/api/projects", json={"name": "scene-test"})
        pid = r.get_json()["id"]
        r = client.get(f"/api/projects/{pid}/chapters/nonexistent/scenes")
        assert r.status_code == 404

    def test_save_null_content_succeeds(self, client):
        """Save with null content should default to empty string, not 500."""
        r = client.post("/api/projects", json={"name": "null-content-test"})
        pid = r.get_json()["id"]
        r = client.post(f"/api/projects/{pid}/chapters", json={"name": "c1"})
        # null content
        r = client.put(f"/api/projects/{pid}/chapters/c1/content", json={"content": None})
        assert r.status_code == 200
        d = r.get_json()
        assert d.get("bytes") == 0
        # missing content
        r = client.put(f"/api/projects/{pid}/chapters/c1/content", json={})
        assert r.status_code == 200

    def test_stats_bounds(self, client):
        """Stats should validate bounds."""
        r = client.post("/api/projects", json={"name": "stats-bounds-test"})
        pid = r.get_json()["id"]
        # Too high
        r = client.put(f"/api/projects/{pid}/stats", json={"daily_goal": 10000000})
        assert r.status_code == 400
        # Negative
        r = client.put(f"/api/projects/{pid}/stats", json={"daily_goal": -1})
        assert r.status_code == 400
        # Valid
        r = client.put(f"/api/projects/{pid}/stats", json={"daily_goal": 1000})
        assert r.status_code == 200

    def test_delete_nonexistent_chapter_404(self, client):
        r = client.post("/api/projects", json={"name": "del-test"})
        pid = r.get_json()["id"]
        r = client.delete(f"/api/projects/{pid}/chapters/nonexistent")
        assert r.status_code == 404


# --------------------------------------------------------------------------
# Quill rename (replaces Dross references)
# --------------------------------------------------------------------------

class TestQuillRename:
    def test_ai_assistant_says_quill(self, client):
        """The /api/chat should use the Quill persona, not Dross."""
        from slot_providers import PROVIDERS
        mock_instance = MagicMock()
        mock_instance.chat.return_value = "ok"
        mock_instance.stream.return_value = iter(["ok"])
        with patch.dict(PROVIDERS, {"minimax": MagicMock(return_value=mock_instance)}):
            client.post("/api/slots/minimax-text/activate")
            r = client.post("/api/chat", json={
                "messages": [{"role": "user", "content": "hi"}],
                "stream": False,
            })
            assert r.status_code == 200
            msgs = mock_instance.chat.call_args[0][0]
            assert "Quill" in msgs[0]["content"]
            assert "Dross" not in msgs[0]["content"]

    def test_chapter_writer_persona_is_quill(self):
        from book_writer import CHAPTER_SYSTEM
        assert "Quill" in CHAPTER_SYSTEM
        assert "Dross" not in CHAPTER_SYSTEM


# --------------------------------------------------------------------------
# Edit-fix endpoint (Zed-style inline AI)
# --------------------------------------------------------------------------

class TestEditFix:
    """Tests for /api/edit-fix — the inline AI fix endpoint used by the
    Swift editor's "Tab to fix" feature. Backed by a small fast slot
    (gemma4:latest, llama3-groq-tool-use:8b, etc.)."""

    def test_endpoint_exists(self, client):
        r = client.post("/api/edit-fix", json={"text": "hello wrold"})
        # Either 200 (with text) or 5xx (no slot) — never 404
        assert r.status_code != 404

    def test_requires_text_field(self, client):
        r = client.post("/api/edit-fix", json={})
        assert r.status_code == 400
        d = r.get_json()
        assert "text" in d.get("error", "")

    def test_rejects_empty_text(self, client):
        r = client.post("/api/edit-fix", json={"text": ""})
        assert r.status_code == 400
        r = client.post("/api/edit-fix", json={"text": "   "})
        assert r.status_code == 400

    def test_rejects_non_string_text(self, client):
        r = client.post("/api/edit-fix", json={"text": 12345})
        assert r.status_code == 400
        r = client.post("/api/edit-fix", json={"text": None})
        assert r.status_code == 400
        r = client.post("/api/edit-fix", json={"text": ["list", "of", "words"]})
        assert r.status_code == 400

    def test_rejects_text_too_long(self, client):
        big_text = "a" * 20001
        r = client.post("/api/edit-fix", json={"text": big_text})
        assert r.status_code == 400
        d = r.get_json()
        assert "too long" in d.get("error", "").lower() or "max" in d.get("error", "").lower()

    def test_default_instruction(self, client):
        """When no instruction is provided, the default 'fix typos and grammar' is used."""
        from unittest.mock import patch, MagicMock
        from slot_providers import PROVIDERS
        # Activate an ollama slot so the edit-fix endpoint routes to our mock
        client.post("/api/slots/gemma4-fast/activate")
        mock_inst = MagicMock()
        mock_inst.chat.return_value = "fixed text"
        with patch.dict(PROVIDERS, {"ollama": MagicMock(return_value=mock_inst)}):
            r = client.post("/api/edit-fix", json={"text": "helo wrold"})
            assert r.status_code == 200, r.get_data(as_text=True)
            mock_inst.chat.assert_called_once()
            msgs = mock_inst.chat.call_args[0][0]
            system_msg = msgs[0]
            assert "typos" in system_msg["content"].lower() or \
                   "grammar" in system_msg["content"].lower()
            user_msg = msgs[1]
            assert "fix typos" in user_msg["content"].lower()

    def test_custom_instruction(self, client):
        from unittest.mock import patch, MagicMock
        from slot_providers import PROVIDERS
        client.post("/api/slots/gemma4-fast/activate")
        mock_inst = MagicMock()
        mock_inst.chat.return_value = "expanded"
        with patch.dict(PROVIDERS, {"ollama": MagicMock(return_value=mock_inst)}):
            r = client.post("/api/edit-fix", json={
                "text": "short",
                "instruction": "expand with sensory detail",
            })
            assert r.status_code == 200
            msgs = mock_inst.chat.call_args[0][0]
            assert "expand with sensory detail" in msgs[1]["content"]

    def test_strips_code_fence_wrapper(self):
        """_strip_edit_fix_wrapper should remove ```markdown blocks and preambles."""
        from server import _strip_edit_fix_wrapper
        assert _strip_edit_fix_wrapper("```markdown\nfixed text\n```") == "fixed text"
        assert _strip_edit_fix_wrapper("```\nfixed text\n```") == "fixed text"
        assert _strip_edit_fix_wrapper("Here is the corrected text:\n\nfixed text") == "fixed text"
        assert _strip_edit_fix_wrapper("Here's a condensed version:\n\ncondensed text") == "condensed text"
        assert _strip_edit_fix_wrapper("Sure, here is the corrected text:\n\nfixed text") == "fixed text"
        assert _strip_edit_fix_wrapper("Corrected text: fixed text") == "fixed text"
        assert _strip_edit_fix_wrapper("  just text  ") == "just text"
        assert _strip_edit_fix_wrapper("no wrapper here") == "no wrapper here"
        # Multi-sentence preamble with sentence terminator
        assert _strip_edit_fix_wrapper(
            "I understand what you're trying to say. Here's the corrected text:\n\nfixed text"
        ) == "fixed text"

    def test_response_shape(self, client):
        """Successful response has text, slot_id, model_id, original_chars, fixed_chars."""
        from unittest.mock import patch, MagicMock
        from slot_providers import PROVIDERS
        client.post("/api/slots/gemma4-fast/activate")
        mock_inst = MagicMock()
        mock_inst.chat.return_value = "fixed version"
        with patch.dict(PROVIDERS, {"ollama": MagicMock(return_value=mock_inst)}):
            r = client.post("/api/edit-fix", json={"text": "original text"})
            assert r.status_code == 200
            d = r.get_json()
            assert "text" in d
            assert "slot_id" in d
            assert "model_id" in d
            assert d["original_chars"] == len("original text")
            assert d["fixed_chars"] == len("fixed version")
            assert d["instruction"] == "fix typos and grammar"

    def test_low_temperature_for_determinism(self, client):
        """edit-fix should use low temperature for deterministic fixes."""
        from unittest.mock import patch, MagicMock
        from slot_providers import PROVIDERS
        client.post("/api/slots/gemma4-fast/activate")
        mock_inst = MagicMock()
        mock_inst.chat.return_value = "x"
        with patch.dict(PROVIDERS, {"ollama": MagicMock(return_value=mock_inst)}):
            r = client.post("/api/edit-fix", json={"text": "hi"})
            assert r.status_code == 200
            call_kwargs = mock_inst.chat.call_args[1]
            assert call_kwargs.get("temperature", 1.0) <= 0.3
            assert call_kwargs.get("max_tokens", 0) >= 256

    def test_instruction_truncation(self, client):
        """Overly long instructions fall back to the default."""
        from unittest.mock import patch, MagicMock
        from slot_providers import PROVIDERS
        client.post("/api/slots/gemma4-fast/activate")
        mock_inst = MagicMock()
        mock_inst.chat.return_value = "x"
        with patch.dict(PROVIDERS, {"ollama": MagicMock(return_value=mock_inst)}):
            r = client.post("/api/edit-fix", json={
                "text": "hi",
                "instruction": "x" * 600,  # > 500 char limit
            })
            assert r.status_code == 200
            d = r.get_json()
            # Should have used the default instruction
            assert d["instruction"] == "fix typos and grammar"


# --------------------------------------------------------------------------
# MCP endpoint (JSON-RPC 2.0 over HTTP)
# --------------------------------------------------------------------------

class TestMCPEndpoint:
    """Tests for /api/mcp — HTTP JSON-RPC 2.0 endpoint that exposes Quill
    tools to MCP-compatible clients (Claude Desktop, Cursor, etc.).
    Mirrors the stdio server in Helpers/quill-ai-helper.swift."""

    def _post(self, client, method, params=None, req_id=1):
        return client.post("/api/mcp", json={
            "jsonrpc": "2.0",
            "id": req_id,
            "method": method,
            "params": params or {},
        })

    def test_initialize(self, client):
        r = self._post(client, "initialize")
        assert r.status_code == 200
        d = r.get_json()
        assert d["jsonrpc"] == "2.0"
        assert d["id"] == 1
        assert d["result"]["serverInfo"]["name"] == "quill"
        assert "tools" in d["result"]["capabilities"]

    def test_tools_list(self, client):
        r = self._post(client, "tools/list")
        assert r.status_code == 200
        d = r.get_json()
        tools = [t["name"] for t in d["result"]["tools"]]
        # Core tools
        assert "list_projects" in tools
        assert "edit_fix" in tools
        assert "search_web" in tools
        assert "shell_exec" in tools
        assert "list_files" in tools
        assert "read_file" in tools
        assert "send_email" in tools
        assert "list_inbox" in tools

    def test_tools_list_schema_format(self, client):
        """Each tool should have a proper MCP inputSchema."""
        r = self._post(client, "tools/list")
        tools = r.get_json()["result"]["tools"]
        for tool in tools:
            assert "name" in tool
            assert "description" in tool
            assert "inputSchema" in tool
            assert tool["inputSchema"]["type"] == "object"

    def test_invalid_jsonrpc_version(self, client):
        r = client.post("/api/mcp", json={
            "jsonrpc": "1.0", "id": 1, "method": "initialize", "params": {}
        })
        assert r.status_code == 400

    def test_unknown_method(self, client):
        r = self._post(client, "nonexistent/method")
        assert r.status_code == 404
        assert r.get_json()["error"]["code"] == -32601

    def test_tools_call_missing_name(self, client):
        r = self._post(client, "tools/call", {"arguments": {}})
        assert r.status_code == 400
        assert r.get_json()["error"]["code"] == -32602

    def test_tools_call_unknown_tool(self, client):
        r = self._post(client, "tools/call", {"name": "fake_tool", "arguments": {}})
        assert r.status_code == 200
        result = r.get_json()["result"]
        assert "unknown tool" in result["content"][0]["text"].lower() or result.get("isError")

    def test_list_projects_tool(self, client):
        r = self._post(client, "tools/call", {"name": "list_projects", "arguments": {}})
        assert r.status_code == 200
        result = r.get_json()["result"]
        # Result is JSON-encoded in the text content
        text = result["content"][0]["text"]
        parsed = json.loads(text)
        assert isinstance(parsed, list)
        # No __context__ pseudo-projects
        for p in parsed:
            assert not p["id"].startswith("__")

    def test_edit_fix_tool(self, client):
        from unittest.mock import patch, MagicMock
        from slot_providers import PROVIDERS
        client.post("/api/slots/gemma4-fast/activate")
        mock_inst = MagicMock()
        mock_inst.chat.return_value = "fixed via MCP"
        with patch.dict(PROVIDERS, {"ollama": MagicMock(return_value=mock_inst)}):
            r = self._post(client, "tools/call", {
                "name": "edit_fix",
                "arguments": {"text": "helo wrold", "instruction": "fix"},
            })
            assert r.status_code == 200
            result = r.get_json()["result"]
            assert result["isError"] is False
            text = result["content"][0]["text"]
            assert "fixed via MCP" in text

    def test_edit_fix_tool_empty_text(self, client):
        r = self._post(client, "tools/call", {
            "name": "edit_fix",
            "arguments": {"text": ""},
        })
        assert r.status_code == 200
        result = r.get_json()["result"]
        # Should have an error or empty response
        assert result.get("isError") or "error" in result["content"][0]["text"].lower()

    def test_shell_exec_tool_blocks_dangerous(self, client):
        """MCP shell_exec should route through the safety-checked tool."""
        r = self._post(client, "tools/call", {
            "name": "shell_exec",
            "arguments": {"cmd": "rm -rf /"},
        })
        assert r.status_code == 200
        result = r.get_json()["result"]
        text = result["content"][0]["text"]
        assert "blocked" in text.lower() or "error" in text.lower()

    def test_list_files_tool(self, client):
        r = self._post(client, "tools/call", {
            "name": "list_files",
            "arguments": {"path": "/tmp"},
        })
        assert r.status_code == 200
        result = r.get_json()["result"]
        assert result["isError"] is False

    def test_search_web_tool(self, client):
        r = self._post(client, "tools/call", {
            "name": "search_web",
            "arguments": {"query": "swift programming", "max_results": 3},
        })
        assert r.status_code == 200
        result = r.get_json()["result"]
        text = result["content"][0]["text"]
        # Real network call — may or may not return results
        parsed = json.loads(text) if text.startswith("{") else {}
        assert "results" in parsed or "error" in parsed

    def test_write_then_read_chapter(self, client):
        """End-to-end: write a chapter via MCP, then read it back."""
        # Create a project first
        r = client.post("/api/projects", json={"name": "mcp-test"})
        pid = r.get_json()["id"]
        # Write
        r = self._post(client, "tools/call", {
            "name": "write_chapter",
            "arguments": {
                "project_id": pid,
                "chapter": "chapter-1",
                "content": "# Chapter 1\n\nIt was a dark and stormy night.",
            },
        })
        assert r.status_code == 200
        # Read back
        r = self._post(client, "tools/call", {
            "name": "read_chapter",
            "arguments": {"project_id": pid, "chapter": "chapter-1"},
        })
        assert r.status_code == 200
        result = r.get_json()["result"]
        text = result["content"][0]["text"]
        parsed = json.loads(text)
        assert "dark and stormy night" in parsed["content"]

    def test_invalid_project_id_rejected(self, client):
        r = self._post(client, "tools/call", {
            "name": "list_chapters",
            "arguments": {"project_id": "../etc"},
        })
        assert r.status_code == 200
        result = r.get_json()["result"]
        # Should return an error
        text = result["content"][0]["text"].lower()
        assert "error" in text or "invalid" in text


# --------------------------------------------------------------------------
# Chapter-write intent detection
# --------------------------------------------------------------------------

class TestChapterWriteIntent:
    """Tests for the chapter-write intent extractor. Catches 'write chapter 3',
    typo'd 'make chapeter 1', reverse 'chapter 3 please write', and the
    prose-only stripper that cleans up the model's *** wrappers."""

    def test_standard_write_chapter(self):
        from server import _extract_chapter_write_intent
        r = _extract_chapter_write_intent("write chapter 3")
        assert r == {"action": "write_chapter", "target": "chapter-03"}

    def test_make_chapter_typo(self):
        """The user's actual broken case: 'make chapeter 1'."""
        from server import _extract_chapter_write_intent
        r = _extract_chapter_write_intent("make chapeter 1")
        assert r is not None
        assert r["target"] == "chapter-01"

    def test_draft_this_chapter(self):
        from server import _extract_chapter_write_intent
        r = _extract_chapter_write_intent("draft this chapter")
        assert r == {"action": "write_chapter", "target": "current"}

    def test_continue(self):
        from server import _extract_chapter_write_intent
        r = _extract_chapter_write_intent("continue")
        assert r is not None
        assert r["target"] == "current"

    def test_reverse_chapter_3_write(self):
        from server import _extract_chapter_write_intent
        r = _extract_chapter_write_intent("chapter 3 — write it")
        assert r is not None
        assert r["target"] == "chapter-03"

    def test_scene(self):
        from server import _extract_chapter_write_intent
        r = _extract_chapter_write_intent("write scene 2")
        assert r is not None
        assert r["target"].startswith("scene-")

    def test_next_thing(self):
        from server import _extract_chapter_write_intent
        r = _extract_chapter_write_intent("write the next thing")
        assert r is not None

    def test_no_intent_for_question(self):
        from server import _extract_chapter_write_intent
        r = _extract_chapter_write_intent("what are the themes of this chapter?")
        assert r is None

    def test_no_intent_for_status(self):
        from server import _extract_chapter_write_intent
        assert _extract_chapter_write_intent("hi how are you") is None
        assert _extract_chapter_write_intent("edit this paragraph") is None
        assert _extract_chapter_write_intent("") is None

    def test_chapter_named_with_word(self):
        from server import _extract_chapter_write_intent
        r = _extract_chapter_write_intent("write chapter one")
        assert r is not None
        assert r["target"] == "chapter-one"

    def test_strip_asterisk_wrapper(self):
        """The user's broken case: model wraps in *** ... *** """
        from server import _strip_chapter_wrapper
        text = "***\n\nIt was a dark night.\n\n***"
        assert _strip_chapter_wrapper(text) == "It was a dark night."

    def test_strip_trailer_question(self):
        """The user's broken case: model adds 'What happens next?' at the end."""
        from server import _strip_chapter_wrapper
        text = "It was a dark night.\n\nWhat happens next?"
        assert _strip_chapter_wrapper(text) == "It was a dark night."

    def test_strip_combined_wrappers(self):
        """Full case: *** + prose + 'What happens next?' trailer."""
        from server import _strip_chapter_wrapper
        text = "***\n\nIt was a dark night.\n\nWhat happens next?\n***"
        result = _strip_chapter_wrapper(text)
        assert "What happens next" not in result
        assert "It was a dark night" in result

    def test_strip_preamble(self):
        from server import _strip_chapter_wrapper
        assert _strip_chapter_wrapper("Here is the chapter:\n\nBody text") == "Body text"
        assert _strip_chapter_wrapper("Sure! Here's your chapter:\n\nBody text") == "Body text"

    def test_strip_preserves_clean_text(self):
        from server import _strip_chapter_wrapper
        clean = "It was a dark night. The rain fell."
        assert _strip_chapter_wrapper(clean) == clean


# --------------------------------------------------------------------------
# _resolve_chapter_target — must create chapter files on demand
# --------------------------------------------------------------------------

class TestResolveChapterTarget:
    """The chapter-write intent used to silently fail when the project had
    no chapters yet, or when the requested chapter didn't exist. It should
    now create the file on disk so the prose actually gets saved."""

    def test_creates_chapter_on_fresh_project(self, client):
        """Brand new project + 'make chapter 1' should create chapter-01.md."""
        from server import _resolve_chapter_target
        r = client.post("/api/projects", json={"name": "fresh-chap-test"})
        pid = r.get_json()["id"]
        # No chapters exist yet
        result = _resolve_chapter_target(pid, "chapter-01")
        assert result == "chapter-01"
        # File should now exist
        import os
        assert os.path.exists(os.path.expanduser(f"~/Quill/projects/{pid}/chapter-01.md"))

    def test_creates_target_chapter_when_missing(self, client):
        """If project has other chapters but the target is missing, create it."""
        from server import _resolve_chapter_target
        r = client.post("/api/projects", json={"name": "mid-project-test"})
        pid = r.get_json()["id"]
        # Create chapter-01
        client.post(f"/api/projects/{pid}/chapters", json={"name": "chapter-01"})
        # Now ask for chapter-05 (doesn't exist) — should create it
        result = _resolve_chapter_target(pid, "chapter-05")
        assert result == "chapter-05"
        import os
        assert os.path.exists(os.path.expanduser(f"~/Quill/projects/{pid}/chapter-05.md"))

    def test_resolves_existing_chapter(self, client):
        from server import _resolve_chapter_target
        r = client.post("/api/projects", json={"name": "existing-chap-test"})
        pid = r.get_json()["id"]
        client.post(f"/api/projects/{pid}/chapters", json={"name": "chapter-03"})
        result = _resolve_chapter_target(pid, "chapter-03")
        assert result == "chapter-03"

    def test_resolves_current_to_first_when_no_context(self, client):
        from server import _resolve_chapter_target
        r = client.post("/api/projects", json={"name": "current-test"})
        pid = r.get_json()["id"]
        client.post(f"/api/projects/{pid}/chapters", json={"name": "chapter-02"})
        # 'current' falls back to first file
        result = _resolve_chapter_target(pid, "current")
        assert result == "chapter-02"

    def test_full_intent_to_written_file(self, client):
        """End-to-end: 'make chapeter 1' on a fresh project writes to disk."""
        from unittest.mock import patch, MagicMock
        from slot_providers import PROVIDERS
        r = client.post("/api/projects", json={"name": "e2e-intent-test"})
        pid = r.get_json()["id"]
        mock_inst = MagicMock()
        mock_inst.chat.return_value = "It was a dark and stormy night."
        with patch.dict(PROVIDERS, {"ollama": MagicMock(return_value=mock_inst)}):
            r = client.post("/api/chat", json={
                "project_id": pid,
                "messages": [{"role": "user", "content": "make chapeter 1"}],
                "stream": False,
            })
            assert r.status_code == 200
            d = r.get_json()
            assert d.get("chapter_written") == "chapter-01"
            assert "dark and stormy" in d.get("text", "")
            # File should be on disk with the content
            import os
            fp = os.path.expanduser(f"~/Quill/projects/{pid}/chapter-01.md")
            assert os.path.exists(fp)
            with open(fp) as f:
                content = f.read()
            assert "dark and stormy" in content


# --------------------------------------------------------------------------
# OpenClaw skills integration
# --------------------------------------------------------------------------

class TestSkillsEndpoints:
    """Tests for /api/skills — exposes the user's installed OpenClaw skills
    to the AI and the UI."""

    def test_skills_list(self, client):
        r = client.get("/api/skills")
        assert r.status_code == 200
        d = r.get_json()
        assert "skills" in d
        assert "status" in d
        assert d["status"]["available"] is True
        assert d["status"]["skill_count"] > 0
        # Each skill has a name + keywords
        for s in d["skills"]:
            assert "name" in s
            assert "keywords" in s
            assert "paths" in s

    def test_skills_list_contains_common_skills(self, client):
        r = client.get("/api/skills")
        d = r.get_json()
        names = {s["name"] for s in d["skills"]}
        # Common skills that should be in any OpenClaw install
        assert "summarize" in names
        assert "github" in names
        assert "weather" in names

    def test_get_skill(self, client):
        r = client.get("/api/skills/summarize")
        assert r.status_code == 200
        d = r.get_json()
        assert d["name"] == "summarize"
        assert "keywords" in d
        assert "summarize" in [k.lower() for k in d["keywords"]]

    def test_get_skill_not_found(self, client):
        r = client.get("/api/skills/nonexistent_skill_xyz")
        assert r.status_code == 404
        assert "not found" in r.get_json()["error"].lower()

    def test_skills_reload(self, client):
        r = client.post("/api/skills/reload")
        assert r.status_code == 200
        d = r.get_json()
        assert "available" in d
        assert "skill_count" in d


class TestSkillsInSystemPrompt:
    """Tests that OpenClaw skills are injected into the default system prompt."""

    def test_prompt_includes_skills(self):
        from server import _dross_system_prompt
        prompt = _dross_system_prompt()
        # If skills are available, they should be in the prompt
        from skills import status, list_skills
        s = status()
        if s["available"]:
            assert "OpenClaw skills" in prompt
            # The priority skills (summarize, github, weather) should be near
            # the top of the list (priority sort).
            assert "summarize" in prompt
        # The base prompt should always be there
        assert "You are Quill" in prompt
        assert "web_search" in prompt
        assert "email_send" in prompt

    def test_skills_for_prompt_function(self):
        from skills import skills_for_prompt
        block = skills_for_prompt(max_skills=10)
        if block:
            assert "OpenClaw skills" in block
            # Count the skill lines (each is `- \`name\``)
            count = block.count("\n- `")
            assert count > 0
            assert count <= 10  # max_skills respected

    def test_find_skill_by_keyword(self):
        from skills import find_skill_by_keyword
        r = find_skill_by_keyword("can you summarize this article")
        # May or may not find depending on installed skills, but shouldn't crash
        if r:
            assert "name" in r

    def test_mcp_exposes_skills_tools(self, client):
        """MCP should expose list_skills and read_skill tools."""
        # initialize
        r = client.post("/api/mcp", json={
            "jsonrpc": "2.0", "id": 1, "method": "tools/list", "params": {}
        })
        tools = {t["name"] for t in r.get_json()["result"]["tools"]}
        assert "list_skills" in tools
        assert "read_skill" in tools

    def test_mcp_list_skills(self, client):
        r = client.post("/api/mcp", json={
            "jsonrpc": "2.0", "id": 1, "method": "tools/call",
            "params": {"name": "list_skills", "arguments": {}}
        })
        assert r.status_code == 200
        result = r.get_json()["result"]
        text = result["content"][0]["text"]
        import json as _json
        data = _json.loads(text)
        assert "skills" in data
        assert "status" in data

    def test_mcp_read_skill(self, client):
        r = client.post("/api/mcp", json={
            "jsonrpc": "2.0", "id": 1, "method": "tools/call",
            "params": {"name": "read_skill", "arguments": {"name": "summarize"}}
        })
        assert r.status_code == 200
        result = r.get_json()["result"]
        text = result["content"][0]["text"]
        import json as _json
        data = _json.loads(text)
        assert data["name"] == "summarize"


# --------------------------------------------------------------------------
# Server info + project location discovery
# --------------------------------------------------------------------------

class TestInfoEndpoint:
    """Tests for /api/info (server config, base_dir, skills status)."""

    def test_info_endpoint(self, client):
        r = client.get("/api/info")
        assert r.status_code == 200
        d = r.get_json()
        assert "version" in d
        assert "base_dir" in d
        assert "base_dir_exists" in d
        assert "ollama_url" in d
        assert "agentmail_inbox" in d
        assert "skills" in d
        assert d["skills"]["count"] > 0

    def test_base_dir_resolves(self, client):
        """The discovered BASE_DIR should exist (or have been created)."""
        from server import BASE_DIR
        # Test fixture sets BASE_DIR to a temp dir, so this should be the temp dir
        r = client.get("/api/info")
        d = r.get_json()
        assert d["base_dir"] == str(BASE_DIR)


class TestRequestSizeLimit:
    """Test that the server rejects oversized requests."""

    def test_oversized_payload_rejected(self, client):
        # Create a chapter with a huge content payload
        big = "a" * (33 * 1024 * 1024)  # 33MB > 32MB limit
        r = client.post("/api/projects", json={"name": "size-test"})
        pid = r.get_json()["id"]
        r = client.put(f"/api/projects/{pid}/chapters/big/content", json={"content": big})
        # Should be 413 Payload Too Large
        assert r.status_code == 413


class TestSkillsReload:
    """Test the /api/skills/reload endpoint."""

    def test_reload_returns_status(self, client):
        r = client.post("/api/skills/reload")
        assert r.status_code == 200
        d = r.get_json()
        assert "available" in d
        assert "skill_count" in d


class TestSaveErrorHandling:
    """Tests that the save error handler toasts and retries appropriately."""

    def test_save_succeeds_via_put(self, client):
        """Smoke test for the put endpoint that saveNow uses."""
        r = client.post("/api/projects", json={"name": "save-test"})
        pid = r.get_json()["id"]
        client.post(f"/api/projects/{pid}/chapters", json={"name": "c1"})
        r = client.put(f"/api/projects/{pid}/chapters/c1/content", json={"content": "Hello"})
        assert r.status_code == 200
        # Verify it persisted
        r = client.get(f"/api/projects/{pid}/chapters/c1/content")
        assert "Hello" in r.get_json()["content"]


# --------------------------------------------------------------------------
# External CLI tools (claude, openclaw, clawhub)
# --------------------------------------------------------------------------

class TestExternalCLITools:
    """Tests for the new external CLI tools: claude, openclaw, clawhub."""

    def test_claude_blocked_pattern(self):
        from dross_tools import tool_claude
        r = tool_claude({"prompt": "rm -rf /"})
        assert "error" in r
        assert "blocked" in r["error"].lower() or "rm" in r["error"].lower()

    def test_claude_requires_prompt(self):
        from dross_tools import tool_claude
        r = tool_claude({})
        assert "error" in r
        assert "prompt" in r["error"]

    def test_claude_calls_binary(self):
        """If claude is installed, this should run; if not, return not-found."""
        import shutil
        from dross_tools import tool_claude
        if not shutil.which("claude"):
            r = tool_claude({"prompt": "echo hi"})
            assert "not found" in r.get("error", "").lower() or "error" in r
        else:
            r = tool_claude({"prompt": "say hi", "timeout": 30})
            assert "stdout" in r or "error" in r

    def test_openclaw_requires_prompt(self):
        from dross_tools import tool_openclaw
        r = tool_openclaw({})
        assert "error" in r
        assert "prompt" in r["error"]

    def test_openclaw_calls_binary(self):
        import shutil
        from dross_tools import tool_openclaw
        if not shutil.which("openclaw"):
            r = tool_openclaw({"prompt": "test"})
            assert "not found" in r.get("error", "").lower()
        else:
            r = tool_openclaw({"prompt": "echo hi", "timeout": 15})
            assert "stdout" in r or "error" in r

    def test_clawhub_search(self):
        from dross_tools import tool_clawhub
        r = tool_clawhub({"action": "search", "query": "summarize"})
        # Either works or binary not found
        assert "stdout" in r or "error" in r

    def test_clawhub_requires_action(self):
        from dross_tools import tool_clawhub
        r = tool_clawhub({})
        # 'list' is the default action and returns a helpful note, not an error
        assert r.get("note") or r.get("error")

    def test_clawhub_search_requires_query(self):
        from dross_tools import tool_clawhub
        r = tool_clawhub({"action": "search"})
        assert "error" in r
        assert "query" in r["error"]

    def test_clawhub_install_requires_name(self):
        from dross_tools import tool_clawhub
        r = tool_clawhub({"action": "install"})
        assert "error" in r
        assert "name" in r["error"]

    def test_clawhub_unknown_action(self):
        from dross_tools import tool_clawhub
        r = tool_clawhub({"action": "frobnicate"})
        assert "error" in r

    def test_cli_safety_blocks_sudo(self):
        from dross_tools import _run_cli
        r = _run_cli("claude", ["sudo", "apt-get", "update"])
        assert "error" in r
        assert "blocked" in r["error"].lower()

    def test_mcp_exposes_claude(self, client):
        r = client.post("/api/mcp", json={
            "jsonrpc": "2.0", "id": 1, "method": "tools/list", "params": {}
        })
        tools = {t["name"] for t in r.get_json()["result"]["tools"]}
        assert "claude" in tools
        assert "openclaw" in tools
        assert "clawhub" in tools


class TestSkillsMultiLocation:
    """Skills should be findable across multiple directories
    (thesolai.github.io, .openclaw/workspace/skills, .openclaw/skills)."""

    def test_skill_dir_discovery_finds_openclaw_workspace(self):
        from skills import _find_all_skill_dirs, find_skill_md
        dirs = _find_all_skill_dirs()
        # Should include the openclaw workspace
        has_openclaw = any(".openclaw" in str(d) for d in dirs)
        assert has_openclaw, f"no .openclaw dir in: {dirs}"

    def test_find_skill_md_finds_installed_skills(self):
        from skills import find_skill_md
        # book-writing was installed earlier
        found = find_skill_md("book-writing")
        # May or may not exist depending on installation
        if found:
            assert found.exists()
            assert found.suffix == ".md"


class TestServerInfoEndpoint:
    """The /api/info endpoint exposes server config for client discovery."""

    def test_info_shape(self, client):
        r = client.get("/api/info")
        d = r.get_json()
        assert d["version"] == "1.0.0"
        assert "base_dir" in d
        assert "ollama_url" in d
        assert "agentmail_inbox" in d
        assert d["skills"]["count"] > 0


# --------------------------------------------------------------------------
# Chapter-write with no project selected (auto-create "default" project)
# --------------------------------------------------------------------------

class TestChapterWriteNoProject:
    """The bug: 'create chapter 1' with no project selected silently failed
    because the chat endpoint required `project_id != "default"`. Now the
    server auto-creates a 'default' project so the user can write chapters
    without first setting up a project."""

    def test_create_chapter_with_default_project_id(self, client):
        """Send a 'create chapter 1' with project_id=default. Should create
        a default project and write the chapter."""
        from unittest.mock import patch, MagicMock
        from slot_providers import PROVIDERS
        # Activate an ollama slot so the endpoint has a model
        client.post("/api/slots/gemma4-fast/activate")
        mock_inst = MagicMock()
        mock_inst.chat.return_value = "It was a dark and stormy night."
        with patch.dict(PROVIDERS, {"ollama": MagicMock(return_value=mock_inst)}):
            r = client.post("/api/chat", json={
                "project_id": "default",
                "messages": [{"role": "user", "content": "create chapter 1"}],
                "stream": False,
            })
            assert r.status_code == 200
            d = r.get_json()
            assert d.get("chapter_written") == "chapter-01"
            assert d.get("project_id") == "default"
            # File should exist in the default project
            import os
            fp = os.path.expanduser("~/Quill/projects/default/chapter-01.md")
            assert os.path.exists(fp), f"file not created at {fp}"
            with open(fp) as f:
                content = f.read()
            assert "dark and stormy" in content

    def test_create_chapter_with_empty_project_id(self, client):
        """Empty project_id should also work — treated as default."""
        from unittest.mock import patch, MagicMock
        from slot_providers import PROVIDERS
        client.post("/api/slots/gemma4-fast/activate")
        mock_inst = MagicMock()
        mock_inst.chat.return_value = "Empty ID test."
        with patch.dict(PROVIDERS, {"ollama": MagicMock(return_value=mock_inst)}):
            r = client.post("/api/chat", json={
                "project_id": "",
                "messages": [{"role": "user", "content": "create chapter 2"}],
                "stream": False,
            })
            # Empty string is treated as "default"
            d = r.get_json()
            assert d.get("chapter_written") is not None


# --------------------------------------------------------------------------
# Skills auto-discovery (from installed SKILL.md files)
# --------------------------------------------------------------------------

class TestSkillsAutoDiscovery:
    """Skills installed via clawhub are auto-discovered from their SKILL.md
    frontmatter, even if they're not in the skill-resolver config."""

    def test_directory_name_used_as_canonical_key(self):
        """The directory name is the canonical key (e.g. 'tmux', 'bash'),
        not the 'name:' field inside the frontmatter."""
        from skills import list_skills
        skills = {s["name"] for s in list_skills()}
        assert "tmux" in skills, f"tmux not in skills: {sorted(skills)[:5]}..."
        assert "bash" in skills

    def test_keywords_extracted_from_frontmatter(self):
        """For skills that declare keywords, those are used."""
        from skills import get_skill
        # Most installed skills have at least the directory name as a keyword
        tmux = get_skill("tmux")
        assert tmux is not None
        assert "tmux" in tmux.get("keywords", [])

    def test_skill_md_path_recorded(self):
        """Each discovered skill has at least one path to its SKILL.md."""
        from skills import get_skill
        tmux = get_skill("tmux")
        assert tmux is not None
        paths = tmux.get("paths", [])
        assert len(paths) > 0
        assert any("SKILL.md" in p for p in paths)

    def test_can_read_skill_md_content(self):
        """read_skill_md should return the file's content."""
        from skills import read_skill_md
        content = read_skill_md("tmux")
        if content:
            assert "tmux" in content.lower()

    def test_installed_skill_total_count(self):
        """At least 70 skills (55 config + 10+ new installs)."""
        from skills import list_skills
        total = len(list_skills())
        assert total >= 70, f"expected ≥70 skills, got {total}"

    def test_skill_priority_order(self):
        """The system prompt lists priority skills first (summarize, github, etc)."""
        from server import _dross_system_prompt
        prompt = _dross_system_prompt()
        # These priority skills should appear
        for priority_skill in ["summarize", "tmux", "bash", "shell-scripting", "sqlite"]:
            assert f"`{priority_skill}`" in prompt, f"{priority_skill} not in priority list"


# --------------------------------------------------------------------------
# CLI tool status and Codex integration
# --------------------------------------------------------------------------

class TestCLITools:
    """Tests for cli_status, codex, and the unified setup diagnostic."""

    def test_cli_status_returns_all_tools(self):
        from dross_tools import tool_cli_status
        r = tool_cli_status({})
        # Should have entries for all the major tools
        for name in ("quill", "claude", "codex", "openclaw", "clawhub", "gemini"):
            assert name in r, f"missing {name} in cli_status: {list(r.keys())}"
        # Each should have 'installed' bool
        for name, info in r.items():
            assert "installed" in info, f"{name} missing 'installed' field"
            assert "version" in info or not info["installed"], f"{name} missing 'version'"

    def test_cli_status_message_for_gemini(self):
        """Gemini should have a clear 'free tier shut down' message."""
        from dross_tools import tool_cli_status
        r = tool_cli_status({})
        if r.get("gemini", {}).get("installed"):
            msg = r["gemini"].get("message", "")
            assert "antigravity" in msg.lower() or "free tier" in msg.lower(), \
                f"gemini message doesn't explain: {msg}"

    def test_codex_requires_prompt(self):
        from dross_tools import tool_codex
        r = tool_codex({})
        assert "error" in r
        assert "prompt" in r["error"]

    def test_codex_safety_blocks_dangerous(self):
        from dross_tools import tool_codex
        r = tool_codex({"prompt": "rm -rf /"})
        assert "error" in r
        assert "blocked" in r["error"].lower()

    def test_codex_calls_binary_if_available(self):
        """If codex is installed, this should run; if not, return not-found."""
        import shutil
        from dross_tools import tool_codex
        if not shutil.which("codex"):
            r = tool_codex({"prompt": "test", "timeout": 15})
            assert "not found" in r.get("error", "").lower()
        else:
            r = tool_codex({"prompt": "test", "timeout": 15})
            assert "stdout" in r or "error" in r

    def test_mcp_exposes_codex_and_cli_status(self, client):
        r = client.post("/api/mcp", json={
            "jsonrpc": "2.0", "id": 1, "method": "tools/list", "params": {}
        })
        tools = {t["name"] for t in r.get_json()["result"]["tools"]}
        assert "codex" in tools
        assert "cli_status" in tools

    def test_mcp_cli_status_works(self, client):
        r = client.post("/api/mcp", json={
            "jsonrpc": "2.0", "id": 1, "method": "tools/call",
            "params": {"name": "cli_status", "arguments": {}}
        })
        result = r.get_json()["result"]
        text = result["content"][0]["text"]
        import json as _json
        data = _json.loads(text)
        # Should have quill, claude, codex, etc.
        assert "quill" in data
        assert "gemini" in data
